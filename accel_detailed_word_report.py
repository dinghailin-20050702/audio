#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate Word report for detailed acceleration insertion analysis."""

import os
import sys
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
import base64

# Configuration
CSV_PATH = os.path.join("analysis_out", "features.csv")
OUTPUT_PATH = os.path.join("analysis_out", "accel_detailed_report.docx")
MAX_DETAILED_PLOTS = 15      # Maximum number of detailed event plots
SCORE_THRESHOLD = 0.5        # Minimum score for detailed analysis
FIG_DPI = 150

def load_data():
    """Load and validate the features CSV file"""
    if not os.path.exists(CSV_PATH):
        print(f"错误: 未找到 {CSV_PATH}")
        print("请先运行 analyze_insertion_audio.py 生成特征文件")
        sys.exit(1)
    
    df = pd.read_csv(CSV_PATH)
    if df.empty:
        print("错误: features.csv 为空")
        sys.exit(1)
    
    return df

def create_summary_table(df):
    """Create summary statistics table data"""
    summary_data = []
    
    # Basic statistics
    summary_data.append(["分析文件总数", str(len(df))])
    
    # Acceleration data availability
    accel_files = df[df['accel_event_score'].notna()]
    summary_data.append(["包含加速度数据的文件", f"{len(accel_files)} ({len(accel_files)/len(df)*100:.1f}%)"])
    
    if len(accel_files) > 0:
        # Score statistics
        score_mean = accel_files['accel_event_score'].mean()
        score_std = accel_files['accel_event_score'].std()
        score_max = accel_files['accel_event_score'].max()
        
        summary_data.extend([
            ["复合评分均值", f"{score_mean:.3f}"],
            ["复合评分标准差", f"{score_std:.3f}"],
            ["复合评分最大值", f"{score_max:.3f}"]
        ])
        
        # High score files
        high_score = accel_files[accel_files['accel_event_score'] > score_mean + score_std]
        summary_data.append(["高评分文件数 (>μ+σ)", f"{len(high_score)} ({len(high_score)/len(accel_files)*100:.1f}%)"])
        
        # Effect sizes
        if 'cohen_d_mag' in accel_files.columns:
            cohen_mean = accel_files['cohen_d_mag'].mean()
            summary_data.append(["平均磁力效应量 (Cohen's d)", f"{cohen_mean:.3f}"])
        
        if 'cohen_d_jerk' in accel_files.columns:
            jerk_mean = accel_files['cohen_d_jerk'].mean()
            summary_data.append(["平均急动效应量 (Cohen's d)", f"{jerk_mean:.3f}"])
    
    return summary_data

def create_ranking_table(df):
    """Create ranking table for top events"""
    # Filter files with acceleration data and sort by score
    accel_files = df[df['accel_event_score'].notna()].copy()
    if len(accel_files) == 0:
        return []
    
    # Sort by composite score
    accel_files_sorted = accel_files.sort_values('accel_event_score', ascending=False)
    
    ranking_data = []
    for i, (_, row) in enumerate(accel_files_sorted.head(20).iterrows(), 1):
        filename = os.path.basename(row['file'])
        score = row['accel_event_score']
        
        # Get key metrics
        mag_peak = row.get('ev_acc_mag_peak', 0)
        jerk_peak = row.get('ev_acc_jerk_peak', 0) 
        cohen_d = row.get('cohen_d_mag', 0)
        p_val = row.get('p_ttest_mag', 1.0)
        
        significance = "***" if p_val < 0.001 else "**" if p_val < 0.01 else "*" if p_val < 0.05 else "ns"
        
        ranking_data.append([
            str(i),
            filename,
            f"{score:.3f}",
            f"{mag_peak:.2f}",
            f"{jerk_peak:.2f}", 
            f"{cohen_d:.3f}",
            significance
        ])
    
    return ranking_data

def create_event_plots(df, max_plots=MAX_DETAILED_PLOTS):
    """Create detailed plots for top events"""
    plots = []
    
    # Get top events
    accel_files = df[df['accel_event_score'].notna()]
    if len(accel_files) == 0:
        return plots
    
    top_events = accel_files.nlargest(max_plots, 'accel_event_score')
    
    for idx, (_, row) in enumerate(top_events.iterrows()):
        try:
            plot_data = create_single_event_plot(row, idx + 1)
            if plot_data:
                plots.append(plot_data)
        except Exception as e:
            print(f"警告: 创建图表失败 {row['file']}: {e}")
    
    return plots

def create_single_event_plot(row, rank):
    """Create a detailed plot for a single event"""
    fig, axes = plt.subplots(2, 2, figsize=(12, 8))
    fig.suptitle(f"#{rank}: {os.path.basename(row['file'])}", fontsize=14, fontweight='bold')
    
    # Plot 1: Acceleration magnitude comparison
    ax1 = axes[0, 0]
    categories = ['事件', '背景']
    ev_mag = row.get('ev_acc_mag_rms', 0)
    bg_mag = row.get('bg_acc_mag_rms', 0)
    
    bars1 = ax1.bar(categories, [ev_mag, bg_mag], color=['red', 'blue'], alpha=0.7)
    ax1.set_ylabel('加速度幅值 RMS')
    ax1.set_title('加速度幅值对比')
    ax1.grid(True, alpha=0.3)
    
    # Add value labels
    for bar, val in zip(bars1, [ev_mag, bg_mag]):
        height = bar.get_height()
        ax1.text(bar.get_x() + bar.get_width()/2., height,
                f'{val:.3f}', ha='center', va='bottom')
    
    # Plot 2: Jerk comparison  
    ax2 = axes[0, 1]
    ev_jerk = row.get('ev_acc_jerk_rms', 0)
    bg_jerk = row.get('bg_acc_jerk_rms', 0)
    
    bars2 = ax2.bar(categories, [ev_jerk, bg_jerk], color=['orange', 'cyan'], alpha=0.7)
    ax2.set_ylabel('急动 RMS')
    ax2.set_title('急动对比')
    ax2.grid(True, alpha=0.3)
    
    for bar, val in zip(bars2, [ev_jerk, bg_jerk]):
        height = bar.get_height()
        ax2.text(bar.get_x() + bar.get_width()/2., height,
                f'{val:.3f}', ha='center', va='bottom')
    
    # Plot 3: PSD band comparison
    ax3 = axes[1, 0]
    bands = ['0-5Hz', '5-10Hz', '10-20Hz', '20-50Hz']
    ev_bands = [row.get(f'ev_acc_psd_band_{b}_ratio', 0) for b in ['0-5', '5-10', '10-20', '20-50']]
    bg_bands = [row.get(f'bg_acc_psd_band_{b}_ratio', 0) for b in ['0-5', '5-10', '10-20', '20-50']]
    
    x = np.arange(len(bands))
    width = 0.35
    
    bars3 = ax3.bar(x - width/2, ev_bands, width, label='事件', color='red', alpha=0.7)
    bars4 = ax3.bar(x + width/2, bg_bands, width, label='背景', color='blue', alpha=0.7)
    
    ax3.set_ylabel('PSD 能量比例')
    ax3.set_title('频域能量分布')
    ax3.set_xticks(x)
    ax3.set_xticklabels(bands)
    ax3.legend()
    ax3.grid(True, alpha=0.3)
    
    # Plot 4: Effect sizes and metrics
    ax4 = axes[1, 1]
    metrics = ['Cohen d (Mag)', 'Cohen d (Jerk)', '复合评分']
    values = [
        row.get('cohen_d_mag', 0),
        row.get('cohen_d_jerk', 0), 
        row.get('accel_event_score', 0)
    ]
    colors = ['green', 'purple', 'gold']
    
    bars5 = ax4.bar(metrics, values, color=colors, alpha=0.7)
    ax4.set_ylabel('效应量 / 评分')
    ax4.set_title('统计效应量')
    ax4.grid(True, alpha=0.3)
    plt.setp(ax4.get_xticklabels(), rotation=45, ha='right')
    
    for bar, val in zip(bars5, values):
        height = bar.get_height()
        ax4.text(bar.get_x() + bar.get_width()/2., height,
                f'{val:.3f}', ha='center', va='bottom')
    
    plt.tight_layout()
    
    # Save plot to memory
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=FIG_DPI, bbox_inches='tight')
    img_buffer.seek(0)
    plt.close()
    
    return {
        'rank': rank,
        'filename': os.path.basename(row['file']),
        'score': row.get('accel_event_score', 0),
        'cohen_d_mag': row.get('cohen_d_mag', 0),
        'cohen_d_jerk': row.get('cohen_d_jerk', 0),
        'p_ttest_mag': row.get('p_ttest_mag', 1.0),
        'image_data': img_buffer.getvalue()
    }

def generate_word_report(df):
    """Generate the complete Word report"""
    doc = Document()
    
    # Title
    title = doc.add_heading('音频插接事件加速度深度分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with timestamp
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'生成时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
    run.font.size = Pt(12)
    run.italic = True
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('1. 执行摘要', level=1)
    
    summary_data = create_summary_table(df)
    summary_table = doc.add_table(rows=len(summary_data), cols=2)
    summary_table.style = 'Table Grid'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, (key, value) in enumerate(summary_data):
        summary_table.cell(i, 0).text = key
        summary_table.cell(i, 1).text = value
    
    # Analysis methodology
    doc.add_heading('2. 分析方法', level=1)
    methodology_text = """
本报告基于音频事件检测与加速度传感器数据的同步分析，采用以下关键技术：

• 音频事件检测：基于RMS能量和谱通量的z-score联合算法
• 加速度特征提取：幅值、急动度（Jerk）的时域与频域分析  
• 统计效应量：Cohen's d测量事件与背景的差异程度
• Welch PSD分析：0-50Hz频段的功率谱密度分布
• 复合评分：多维特征加权组合的事件强度评估

复合评分权重配置：
- 加速度幅值差异(diff_acc_mag_rms): 30%
- 急动度差异(diff_acc_jerk_rms): 20%  
- 幅值效应量(cohen_d_mag): 25%
- 急动效应量(cohen_d_jerk): 15%
- 音频RMS差异(diff_rms_dbfs): 10%
    """
    doc.add_paragraph(methodology_text)
    
    # Event ranking table
    doc.add_heading('3. 事件排名', level=1)
    doc.add_paragraph('基于复合评分的前20个插接事件：')
    
    ranking_data = create_ranking_table(df)
    if ranking_data:
        ranking_table = doc.add_table(rows=len(ranking_data) + 1, cols=7)
        ranking_table.style = 'Table Grid'
        
        # Header
        headers = ['排名', '文件名', '复合评分', '峰值加速度', '峰值急动', "Cohen's d", '显著性']
        for i, header in enumerate(headers):
            cell = ranking_table.cell(0, i)
            cell.text = header
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
        
        # Data rows
        for row_idx, row_data in enumerate(ranking_data, 1):
            for col_idx, cell_data in enumerate(row_data):
                ranking_table.cell(row_idx, col_idx).text = cell_data
    
    # Detailed event analysis
    doc.add_heading('4. 详细事件分析', level=1)
    doc.add_paragraph(f'以下展示复合评分最高的前{MAX_DETAILED_PLOTS}个事件的详细分析图表：')
    
    plots = create_event_plots(df, MAX_DETAILED_PLOTS)
    
    for plot_data in plots:
        # Event section header
        doc.add_heading(f"事件 #{plot_data['rank']}: {plot_data['filename']}", level=2)
        
        # Key metrics paragraph
        metrics_text = f"""
复合评分: {plot_data['score']:.3f}
Cohen's d (幅值): {plot_data['cohen_d_mag']:.3f}  
Cohen's d (急动): {plot_data['cohen_d_jerk']:.3f}
统计显著性 (p值): {plot_data['p_ttest_mag']:.4f}
        """
        doc.add_paragraph(metrics_text)
        
        # Add plot image
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        
        # Create temporary image file
        temp_img_path = f"/tmp/temp_plot_{plot_data['rank']}.png"
        with open(temp_img_path, 'wb') as f:
            f.write(plot_data['image_data'])
        
        run.add_picture(temp_img_path, width=Inches(6))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Clean up temp file
        os.remove(temp_img_path)
        
        doc.add_page_break()
    
    # Conclusions
    doc.add_heading('5. 结论与建议', level=1)
    
    accel_files = df[df['accel_event_score'].notna()]
    if len(accel_files) > 0:
        conclusions_text = f"""
基于{len(accel_files)}个包含加速度数据的文件分析：

关键发现：
• 平均复合评分: {accel_files['accel_event_score'].mean():.3f}
• 最高复合评分: {accel_files['accel_event_score'].max():.3f}  
• 具有统计显著性的事件: {len(accel_files[accel_files['p_ttest_mag'] < 0.05])} 个

建议：
1. 重点关注复合评分 > {accel_files['accel_event_score'].mean() + accel_files['accel_event_score'].std():.3f} 的事件
2. 结合Cohen's d效应量进行进一步验证
3. 考虑调整传感器位置或采样频率以提高检测精度
4. 建立阈值体系用于实时监测应用
        """
    else:
        conclusions_text = "未检测到有效的加速度数据，建议检查传感器配置和数据采集流程。"
    
    doc.add_paragraph(conclusions_text)
    
    return doc

def main():
    """Main function"""
    print("正在加载数据...")
    df = load_data()
    
    print("正在生成 Word 报告...")
    doc = generate_word_report(df)
    
    # Ensure output directory exists
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    
    print(f"正在保存报告到 {OUTPUT_PATH}...")
    doc.save(OUTPUT_PATH)
    
    print(f"报告生成完成！")
    print(f"输出文件: {OUTPUT_PATH}")
    
    # Print summary
    accel_files = df[df['accel_event_score'].notna()]
    print(f"\n摘要:")
    print(f"- 总文件数: {len(df)}")
    print(f"- 包含加速度数据: {len(accel_files)}")
    if len(accel_files) > 0:
        print(f"- 平均复合评分: {accel_files['accel_event_score'].mean():.3f}")
        print(f"- 最高复合评分: {accel_files['accel_event_score'].max():.3f}")

if __name__ == "__main__":
    main()